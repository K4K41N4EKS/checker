from docx import Document
from lxml import etree
from zipfile import ZipFile
from src.python.utils.logger_utils import get_logger

logger = get_logger("toc_checker")


def contains_marker_raw(doc_path: str, marker="Оглавление") -> bool:
    try:
        with ZipFile(doc_path) as docx:
            with docx.open("word/document.xml") as xml_file:
                tree = etree.parse(xml_file)
                text_elements = tree.xpath(
                    "//w:t",
                    namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                )
                for i, el in enumerate(text_elements):
                    if el.text and marker.lower() in el.text.lower():
                        logger.info(f"[FOUND IN XML] '{marker}' найден в теге w:t #{i}: {el.text}")
                        return True
    except Exception as e:
        logger.warning(f"[XML PARSE ERROR] Не удалось открыть DOCX: {e}")
    return False


def find_start_index(doc: Document, marker: str = "Оглавление", doc_path: str = "") -> int | None:
    marker = marker.strip().lower()

    # 1. Поиск в параграфах напрямую
    for i, para in enumerate(doc.paragraphs):
        if marker in para.text.strip().lower():
            logger.info(f"[FOUND IN PARAGRAPH] #{i}: '{para.text.strip()}'")
            return i + 1

    # 2. Поиск в XML
    if doc_path and contains_marker_raw(doc_path, marker):
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            if text.startswith("1 "):
                logger.warning(f"[FALLBACK TO SECTION] '{marker}' найден в XML, fallback на #{i}: '{para.text.strip()}'")
                return i

        for i, para in enumerate(doc.paragraphs):
            if para.text and len(para.text.strip()) > 10:
                logger.warning(f"[FALLBACK TO TEXT] '{marker}' найден в XML, fallback на #{i}: '{para.text.strip()}'")
                return i

        logger.warning(f"[FOUND IN RAW] '{marker}' найден в XML, но fallback не дал результата")
        return None

    logger.warning(f"[NOT FOUND] '{marker}' не найден ни в тексте, ни в XML")
    return None
