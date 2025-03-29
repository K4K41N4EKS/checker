import zipfile
from pathlib import Path
from docx import Document
from lxml import etree
from docx.oxml.ns import qn
from tempfile import TemporaryDirectory
from src.python.utils.docx_zip_comments import add_comment_to_docx


def extract_commented_text(docx_path):
    """
    Извлекает текст из документа, который попал под выделение комментария.
    Возвращает словарь {id_комментария: текст}.
    """
    with zipfile.ZipFile(docx_path, 'r') as z:
        xml = z.read('word/document.xml')
    root = etree.fromstring(xml)
    result = {}
    current_id = None
    text_acc = ""

    for elem in root.iter():
        if elem.tag.endswith('commentRangeStart'):
            current_id = elem.attrib.get(qn('w:id'))
            text_acc = ""
        elif elem.tag.endswith('t') and current_id:
            text_acc += elem.text or ''
        elif elem.tag.endswith('commentRangeEnd') and current_id:
            result[current_id] = text_acc
            current_id = None

    return result


def test_complex_comment_cases():
    """
    Проверяет добавление комментариев к нескольким абзацам и повторное комментирование.
    Убеждается, что выделенный текст соответствует ожидаемому.
    """
    with TemporaryDirectory() as tmp_dir:
        input_docx = Path(tmp_dir) / "test_complex_comments.docx"
        doc = Document()
        doc.add_paragraph("Это первый абзац для теста.")                             # idx 0
        doc.add_paragraph("Второй абзац содержит интересные слова для анализа.")    # idx 1
        doc.add_paragraph("Третий абзац используется повторно.")                    # idx 2
        doc.add_paragraph("Последний абзац без комментариев.")                      # idx 3
        doc.save(input_docx)

        path = str(input_docx)

        path = add_comment_to_docx(path, 0, "Комментарий ко всему первому абзацу.")
        path = add_comment_to_docx(path, 0, "Комментарий к слову 'теста'", word_index=4)
        path = add_comment_to_docx(path, 1, "Комментарий к слову 'интересные'", word_index=3)
        path = add_comment_to_docx(path, 1, "Комментарий ко второму абзацу полностью.")
        path = add_comment_to_docx(path, 2, "Комментарий к слову 'повторно'", word_index=3)
        path = add_comment_to_docx(path, 2, "Общий комментарий к третьему абзацу.")
        path = add_comment_to_docx(path, 2, "Общий комментарий к третьему абзацу.")

        paragraphs = [p.text for p in Document(path).paragraphs]
        expected_paragraphs = [
            "Это первый абзац для теста.",
            "Второй абзац содержит интересные слова для анализа.",
            "Третий абзац используется повторно.",
            "Последний абзац без комментариев.",
        ]
        for expected in expected_paragraphs:
            assert expected in paragraphs, f"Абзац отсутствует: {expected}"

        highlights = extract_commented_text(path)
        assert any("первый абзац" in t for t in highlights.values())
        assert any("интересные слова" in t or "второй абзац" in t for t in highlights.values())
        assert any("повторно" in t or "третий абзац" in t for t in highlights.values())

        print("Все абзацы прокомментированы корректно.")
